
# import pandas as pd
from docx.api import Document
import glob
import numpy as np
import itertools
from collections import OrderedDict


def bubble_sort(array):
    print('---bubble sort---')
    n = len(array)

    for i in range(n):
        # Create a flag that will allow the function to
        # terminate early if there's nothing left to sort
        already_sorted = True

        # Start looking at each item of the list one by one,
        # comparing it with its adjacent value. With each
        # iteration, the portion of the array that you look at
        # shrinks because the remaining items have already been
        # sorted.
        for j in range(n - i - 1):
            print('array j', array[j])
            if array[j][3][1] > array[j + 1][3][1]:
                # If the item you're looking at is greater than its
                # adjacent value, then swap them
                temp = list(array[j][3][1])
                array[j][3][1] = array[j + 1][3][1]
                array[j + 1][3][1] = tuple(temp)

                # array[j][3][1], array[j + 1][3][1] = array[j + 1][3][1], array[j][3][1]



                # Since you had to swap two elements,
                # set the `already_sorted` flag to `False` so the
                # algorithm doesn't finish prematurely
                already_sorted = False

        # If there were no swaps during the last iteration,
        # the array is already sorted, and you can terminate
        if already_sorted:
            break

    return array





baseDir = 'C:\\Users\\KyleRid\\Desktop\\py\\source\\docx' # Starting directory for directory walk

docFiles = glob.glob(baseDir+'/doc*.docx') # for the loop

header = ()
data = []
counter = 0

for i in docFiles:
    counter = 0
    print(i)
    document = Document(i)
    table = document.tables[0]


    # for row in table.rows:
    #     for cell in row.cells:
    #         for para in cell.paragraphs:
    #             print(para.text)


    keys = None
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)
        if i == 0:
            keys = tuple(text)
            continue
        counter = counter + 1
        row_data = list(zip(keys, text))
        if not header and counter == 1:
            header = row_data
        if counter == 2:
            data.append(row_data)
        print(counter)


# print(type(data), data[0], type(data[0]))
doubles = []

for i in data:
    if i not in doubles:
        doubles.append(i)
    else:
        print(i,end=' ')
# print('end', len(doubles))
# i = 0
# while i < len(data):
#     j = i + 1
#     print('i:', i)
#     while j < len(data):
#         print('j:', j)
#         if data[i] == data[j]:
#             doubles.append(data[i])
#             break
#         j += 1
#     i += 1

# print(len(doubles))
# for item in data: #row
#     print(type(item))
#     for i in data:
#         if (item == i) {
#             doubles
#         }
    # for el in item: #cell
        # print(el)

targetDoc = Document('C:\\Users\\KyleRid\\Desktop\\py\\source\\target.docx')
targetTable = targetDoc.tables[0]
globalCounter = 1
# # # add a data row for each item

# newData = bubble_sort(data)
data.sort(key = lambda x: x[3])

# res = list(set([ele for ele in data if data.count(ele) > 1]))


# print(res)


# for item in data:

print(len(data))
# print(type(newData))
for item in data:
    # print(item[3][1])
    print(globalCounter)
    cells = targetTable.add_row().cells
    cells[0].text = str(globalCounter)
    cells[1].text = item[1][1]
    cells[2].text = item[2][1]
    cells[3].text = item[3][1]
    cells[4].text = item[4][1]
    cells[5].text = item[5][1]
    cells[6].text = item[6][1]
    cells[7].text = item[7][1]
    cells[8].text = item[8][1]
    cells[9].text = item[9][1]
    # print(cells[8].text)
    globalCounter = globalCounter + 1

targetDoc.save('C:\\Users\\KyleRid\\Desktop\\py\\source\\target2.docx')











#copy
# header = ()
# data = []
# counter = 0
# keys = None
# for i, row in enumerate(table.rows):
#     text = (cell.text for cell in row.cells)
#     if i == 0:
#         keys = tuple(text)
#         continue
#     counter = counter + 1
#     row_data = dict(zip(keys, text))
#     if not header and counter == 1:
#         header = row_data
#     if counter == 2:
#         data.append(row_data)

# print(data)










# print (header)

# df = pd.DataFrame(data)
